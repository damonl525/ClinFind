import os
import sys
import shutil
import time
from docx import Document

# Add backend to path
BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BACKEND_DIR)
print(f"Added {BACKEND_DIR} to sys.path")

from app.core.database import init_db, get_db_connection
from app.services.indexer import Indexer
from app.services.search_engine import SearchEngine

TEST_DIR = os.path.join(os.path.dirname(__file__), "test_data")

def create_test_files():
    if os.path.exists(TEST_DIR):
        shutil.rmtree(TEST_DIR)
    os.makedirs(TEST_DIR)

    # 1. Create a text file
    with open(os.path.join(TEST_DIR, "sample.txt"), "w", encoding="utf-8") as f:
        f.write("This is a sample text file for FileSearcher verification.\nIt contains the keyword: Banana.")

    # 2. Create a Python file
    with open(os.path.join(TEST_DIR, "script.py"), "w", encoding="utf-8") as f:
        f.write("def calculate_sum(a, b):\n    return a + b\n\nclass DataProcessor:\n    pass")

    # 3. Create a Docx file
    doc = Document()
    doc.add_heading('Project Report', 0)
    doc.add_paragraph('The revenue for Q1 was amazing. We sold 1000 units.')
    doc.save(os.path.join(TEST_DIR, "report.docx"))
    
    print(f"Created test files in: {TEST_DIR}")

def run_verification():
    print("--- Starting Verification ---")
    
    # Initialize DB
    init_db()
    
    # Create files
    create_test_files()
    
    # Run Indexing
    print("\n--- Indexing Folder ---")
    indexer = Indexer()
    indexer.index_folder(TEST_DIR)
    
    # Run Search
    engine = SearchEngine()
    
    queries = ["Banana", "calculate_sum", "revenue"]
    
    print("\n--- Running Search Tests ---")
    for q in queries:
        print(f"\nSearching for: '{q}'")
        results = engine.search(q)
        if results:
            for r in results:
                print(f"  [MATCH] {os.path.basename(r['file_path'])}")
                print(f"  Snippet: {r['highlight']}")
        else:
            print("  [NO MATCH] Found nothing.")

if __name__ == "__main__":
    run_verification()
